"""Document text extraction with scanned-PDF OCR routing.

PDF routing (mode auto):
1. pypdf probe decides whether the PDF has useful embedded text.
2. Embedded text → docling-serve (no OCR) when configured, else local pypdf.
3. Scanned → configured OCR engine; `none` rejects with a clear message.

Never claim a scanned document was anonymized when no text was extracted.
"""

import io
import re
import zipfile

from pydantic import BaseModel, Field

from ..core.config import Settings
from ..services.docling_serve_client import DoclingServeClient, DoclingServeError
from ..services.pdf_text_probe import has_embedded_text


class ExtractionError(Exception):
    def __init__(self, message: str, status_code: int = 422):
        super().__init__(message)
        self.status_code = status_code


class PageRange(BaseModel):
    page_number: int
    start: int
    end: int


class LayoutLine(BaseModel):
    """A transcribed line with its normalized (0–1000, top-left origin)
    bounding box and its character range in the extracted text. Basis for the
    layout-preserving redacted-PDF reconstruction of scanned documents."""

    page_number: int
    x1: int
    y1: int
    x2: int
    y2: int
    start: int
    end: int


class ExtractedDocument(BaseModel):
    text: str
    source_type: str  # "txt" | "docx" | "pdf" | "pdf-ocr" | "paste"
    pages: list[PageRange] = Field(default_factory=list)
    layout: list[LayoutLine] = Field(default_factory=list)
    warnings: list[str] = Field(default_factory=list)


async def extract_document(data: bytes, filename: str, settings: Settings) -> ExtractedDocument:
    suffix = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if suffix == "txt":
        return extract_txt(data)
    if suffix == "docx":
        return extract_docx(data)
    if suffix == "pdf":
        return await extract_pdf(data, filename, settings)
    raise ExtractionError("Unsupported file type; allowed: .txt, .docx, .pdf", status_code=415)


# --- TXT ---------------------------------------------------------------------


def extract_txt(data: bytes) -> ExtractedDocument:
    if b"\x00" in data:
        raise ExtractionError("The file appears to be binary, not UTF-8 text.", status_code=415)
    try:
        return ExtractedDocument(text=data.decode("utf-8-sig"), source_type="txt")
    except UnicodeDecodeError as exc:
        raise ExtractionError("The file is not valid UTF-8 text.", status_code=415) from exc


# --- DOCX --------------------------------------------------------------------

_DOCX_WARNING_MARKERS = [
    (b"<w:ins ", "The document contains tracked changes; inserted text may be incomplete."),
    (b"<w:del ", "The document contains tracked deletions, which are not extracted."),
    (b"txbxContent", "The document contains text boxes, which are not extracted."),
]


def extract_docx(data: bytes) -> ExtractedDocument:
    import docx
    from docx.document import Document as DocxDocument
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    try:
        document: DocxDocument = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(
            "The DOCX file could not be read (malformed or password-protected).", status_code=415
        ) from exc

    parts: list[str] = []

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)

    # Body blocks in document order (paragraphs interleaved with tables).
    body = document.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            text = Paragraph(child, document).text
            if text.strip():
                parts.append(text)
        elif child.tag.endswith("}tbl"):
            table = Table(child, document)
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

    for section in document.sections:
        for paragraph in section.footer.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)

    warnings = _docx_warnings(data)
    text = "\n".join(parts)
    if not text.strip():
        raise ExtractionError("No text could be extracted from the DOCX file.")
    return ExtractedDocument(text=text, source_type="docx", warnings=warnings)


def _docx_warnings(data: bytes) -> list[str]:
    warnings: list[str] = []
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            names = set(archive.namelist())
            if "word/comments.xml" in names:
                warnings.append("The document contains comments, which are not extracted.")
            document_xml = archive.read("word/document.xml")
            for marker, message in _DOCX_WARNING_MARKERS:
                if marker in document_xml:
                    warnings.append(message)
    except Exception:
        pass
    return warnings


# --- PDF ---------------------------------------------------------------------


async def extract_pdf(data: bytes, filename: str, settings: Settings) -> ExtractedDocument:
    if has_embedded_text(
        data,
        min_chars=settings.DOCLING_MIN_EXTRACTED_CHARS_PDF,
        max_pages_to_check=settings.PDF_MAX_PAGES_FOR_TEXT_PROBE,
    ):
        if settings.DOCLING_SERVE_URL:
            try:
                client = DoclingServeClient(settings.DOCLING_SERVE_URL)
                text = await client.convert_pdf(data, filename, do_ocr=False)
                return ExtractedDocument(
                    text=text,
                    source_type="pdf",
                    warnings=["Page mapping is not available for docling-serve extraction."],
                )
            except DoclingServeError as exc:
                # Fall back to local extraction rather than failing the request.
                fallback = _extract_pdf_local(data)
                fallback.warnings.insert(0, f"docling-serve failed ({exc}); used local extraction.")
                return fallback
        return _extract_pdf_local(data)

    # Likely scanned → OCR routing.
    engine = settings.OCR_ENGINE
    if engine == "none":
        raise ExtractionError(
            "This PDF appears to contain scanned images and no extractable text. "
            "OCR is not enabled (OCR_ENGINE=none)."
        )
    if engine == "docling_tesseract":
        if not settings.DOCLING_SERVE_URL:
            raise ExtractionError(
                "OCR engine 'docling_tesseract' requires DOCLING_SERVE_URL.", status_code=503
            )
        try:
            client = DoclingServeClient(settings.DOCLING_SERVE_URL)
            text = await client.convert_pdf(data, filename, do_ocr=True, force_ocr=True)
        except DoclingServeError as exc:
            raise ExtractionError(f"OCR failed: {exc}", status_code=502) from exc
        if not text.strip():
            raise ExtractionError("OCR produced no text for this document.")
        return ExtractedDocument(
            text=text,
            source_type="pdf-ocr",
            warnings=["Text was produced by OCR; recognition errors are possible."],
        )
    if engine == "llm_vision":
        from ..services.vision_llm_ocr import VisionOCRError, VisionOCRService

        try:
            page_lines = await VisionOCRService(settings).process_pdf(data)
        except VisionOCRError as exc:
            raise ExtractionError(str(exc), status_code=exc.status_code) from exc

        parts: list[str] = []
        pages: list[PageRange] = []
        layout: list[LayoutLine] = []
        offset = 0
        for number, lines in enumerate(page_lines, start=1):
            page_start = offset
            for index, line in enumerate(lines):
                if index > 0:
                    offset += 1  # "\n" between lines
                if line.box is not None:
                    x1, y1, x2, y2 = line.box
                    layout.append(
                        LayoutLine(
                            page_number=number,
                            x1=x1,
                            y1=y1,
                            x2=x2,
                            y2=y2,
                            start=offset,
                            end=offset + len(line.text),
                        )
                    )
                offset += len(line.text)
            parts.append("\n".join(line.text for line in lines))
            pages.append(PageRange(page_number=number, start=page_start, end=offset))
            offset += 2  # "\n\n" between pages
        text = "\n\n".join(parts)
        if not text.strip():
            raise ExtractionError("OCR produced no text for this document.")
        return ExtractedDocument(
            text=text,
            source_type="pdf-ocr",
            pages=pages,
            layout=layout,
            warnings=["Text was produced by OCR; recognition errors are possible."],
        )
    if engine == "mistral_ocr":
        raise ExtractionError(
            "OCR engine 'mistral_ocr' is not available yet in this milestone.", status_code=501
        )
    raise ExtractionError(f"Unknown OCR engine '{engine}'.", status_code=500)


def _extract_pdf_local(data: bytes) -> ExtractedDocument:
    from pypdf import PdfReader
    from pypdf.errors import PyPdfError

    try:
        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            if not reader.decrypt(""):
                raise ExtractionError(
                    "The PDF is password-protected and cannot be processed.", status_code=415
                )
    except ExtractionError:
        raise
    except PyPdfError as exc:
        raise ExtractionError(
            "The PDF file could not be read (malformed).", status_code=415
        ) from exc

    parts: list[str] = []
    pages: list[PageRange] = []
    offset = 0
    for number, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        page_text = _normalize_pdf_text(page_text)
        parts.append(page_text)
        pages.append(PageRange(page_number=number, start=offset, end=offset + len(page_text)))
        offset += len(page_text) + 2  # "\n\n" page separator

    text = "\n\n".join(parts)
    if not text.strip():
        raise ExtractionError("No text could be extracted from the PDF.")
    return ExtractedDocument(text=text, source_type="pdf", pages=pages)


def _normalize_pdf_text(text: str) -> str:
    # pypdf occasionally emits stray carriage returns and NULs.
    return re.sub(r"\r\n?", "\n", text).replace("\x00", "")
