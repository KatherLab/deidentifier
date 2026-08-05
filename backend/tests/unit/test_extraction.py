import io

import pytest

from backend.src.core.config import Settings
from backend.src.utils.extraction import (
    ExtractionError,
    extract_document,
    extract_docx,
    extract_txt,
)
from backend.tests.pdf_builder import make_pdf, make_scanned_pdf


def settings_with(**kwargs) -> Settings:
    return Settings(**kwargs)


# --- TXT ---------------------------------------------------------------------


def test_txt_utf8_with_bom():
    doc = extract_txt("﻿Ärztlicher Befund".encode())
    assert doc.text == "Ärztlicher Befund"
    assert doc.source_type == "txt"


def test_txt_binary_rejected():
    with pytest.raises(ExtractionError):
        extract_txt(b"\x00\x01\x02")


# --- DOCX --------------------------------------------------------------------


def make_docx() -> bytes:
    import docx

    document = docx.Document()
    document.add_paragraph("Patient: Max Mustermann")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Fallnummer"
    table.rows[0].cells[1].text = "2024-004711"
    document.add_paragraph("Aufnahme am 10.03.2024")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_docx_paragraphs_and_tables_in_order():
    doc = extract_docx(make_docx())
    assert doc.source_type == "docx"
    assert "Patient: Max Mustermann" in doc.text
    assert "Fallnummer | 2024-004711" in doc.text
    assert doc.text.index("Max Mustermann") < doc.text.index("Fallnummer")
    assert doc.text.index("Fallnummer") < doc.text.index("10.03.2024")


def test_docx_malformed_rejected():
    with pytest.raises(ExtractionError):
        extract_docx(b"not a zip archive")


# --- PDF ---------------------------------------------------------------------


async def test_pdf_text_extracted_locally_without_docling():
    pdf = make_pdf(
        [
            "Patient: Max Mustermann, geb. 01.02.1980",
            "Tel.: 0351 458-0",
            "Der Patient wurde stationaer aufgenommen und komplikationslos behandelt.",
            "Die Entlassung erfolgte in gutem Allgemeinzustand nach Hause.",
        ]
    )
    doc = await extract_document(pdf, "brief.pdf", settings_with())
    assert doc.source_type == "pdf"
    assert "Max Mustermann" in doc.text
    assert "0351 458-0" in doc.text
    assert doc.pages and doc.pages[0].page_number == 1
    # Page offsets must map back into the text.
    page = doc.pages[0]
    assert "Max Mustermann" in doc.text[page.start : page.end]


async def test_scanned_pdf_rejected_when_ocr_disabled():
    pdf = make_scanned_pdf(pages=2)
    with pytest.raises(ExtractionError) as excinfo:
        await extract_document(pdf, "scan.pdf", settings_with(OCR_ENGINE="none"))
    assert "scanned" in str(excinfo.value)


async def test_scanned_pdf_with_unavailable_engine():
    pdf = make_scanned_pdf()
    with pytest.raises(ExtractionError) as excinfo:
        await extract_document(pdf, "scan.pdf", settings_with(OCR_ENGINE="mistral_ocr"))
    assert excinfo.value.status_code == 501


async def test_force_ocr_bypasses_text_probe():
    # A PDF WITH a good embedded text layer would normally extract locally as
    # "pdf"; force_ocr must skip the probe and route to OCR instead. With no OCR
    # engine that surfaces as the dedicated "OCR was requested" error (503) —
    # proving the embedded text was deliberately ignored.
    pdf = make_pdf(["Patient: Max Mustermann, geb. 01.02.1980"])
    with pytest.raises(ExtractionError) as excinfo:
        await extract_document(
            pdf, "brief.pdf", settings_with(OCR_ENGINE="none"), force_ocr=True
        )
    assert excinfo.value.status_code == 503
    assert "OCR was requested" in str(excinfo.value)


async def test_force_ocr_routes_text_pdf_through_ocr_engine(monkeypatch):
    # With an OCR engine available, force_ocr re-OCRs the text PDF: the result
    # is the OCR output (source_type "pdf-ocr"), not the embedded text layer.
    from backend.src.services import docling_serve_client

    async def fake_convert(self, data, filename, *, do_ocr=False, force_ocr=False):
        assert do_ocr and force_ocr
        return "OCR: Herr Mustermann"

    monkeypatch.setattr(docling_serve_client.DoclingServeClient, "convert_pdf", fake_convert)
    pdf = make_pdf(["Patient: Max Mustermann, geb. 01.02.1980"])
    doc = await extract_document(
        pdf,
        "brief.pdf",
        settings_with(OCR_ENGINE="docling_tesseract", DOCLING_SERVE_URL="http://docling"),
        force_ocr=True,
    )
    assert doc.source_type == "pdf-ocr"
    assert doc.text == "OCR: Herr Mustermann"


async def test_unknown_extension_rejected():
    with pytest.raises(ExtractionError) as excinfo:
        await extract_document(b"data", "tabelle.xlsx", settings_with())
    assert excinfo.value.status_code == 415
